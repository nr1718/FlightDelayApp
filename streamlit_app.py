# streamlit_app.py
import os
import uuid
import numpy as np
import pandas as pd
import streamlit as st
import plotly.express as px
import plotly.graph_objects as go
from sklearn.preprocessing import LabelEncoder, StandardScaler
from sklearn.model_selection import train_test_split
from sklearn.linear_model import LinearRegression
import joblib
from sklearn.metrics import mean_absolute_error, r2_score # CRITICAL: Import for metrics
from docx import Document
from docx.shared import Inches
from io import BytesIO
from datetime import datetime, timedelta

# ------------------- Prepare folders -------------------
os.makedirs('models', exist_ok=True)
os.makedirs('outputs', exist_ok=True)

# ------------------- Streamlit Config -------------------
st.set_page_config(
    page_title="Flight Delay Prediction Dashboard",
    layout="wide",
    initial_sidebar_state="expanded"
)

st.title("✈️ Flight Delay Prediction Dashboard")
st.caption("Developed by Nirmal Raj | Flight Delay App © 2025") 

# ------------------- Load or Generate Data -------------------
DATA_FILE = 'flight_delay_data.csv'
MODEL_FILE = 'models/best_model.pkl'

# Define the dictionary for renaming columns (Old Name: New Name)
COLUMN_RENAMES = {
    "destination": "dest",
    "sched_dep_hour": "departure_time",
    "sched_arr_hour": "arrival_time",
    "delay_minutes": "delay"
}

if os.path.exists(DATA_FILE):
    df = pd.read_csv(DATA_FILE)
    
    # Rename columns if they exist (for existing CSVs)
    columns_to_rename = {k: v for k, v in COLUMN_RENAMES.items() if k in df.columns}
    df.rename(columns=columns_to_rename, inplace=True)
    
else:
    st.warning("⚠️ Data not found. Generating synthetic dataset...")
    np.random.seed(42)
    airlines = ['AirA', 'AirB', 'AirC', 'AirD', 'AirE']
    airports = ['BLR','DEL','BOM','MAA','COK','HYD','CCU','AMD']
    weather_conditions = ['Clear','Rain','Storm','Fog','Snow','Windy']

    rows=[]
    for i in range(1000):
        airline=np.random.choice(airlines)
        origin=np.random.choice(airports)
        dest=np.random.choice([a for a in airports if a!=origin])
        dep_hour=np.random.randint(0,24)
        arr_hour=(dep_hour+np.random.randint(1,5))%24
        distance=np.random.randint(100,2000)
        day=np.random.randint(0,7)
        month=np.random.randint(1,13)
        weather=np.random.choice(weather_conditions,p=[0.6,0.2,0.05,0.08,0.01,0.06])
        delay=max(0,int(np.random.normal(loc=10 + 0.02*distance + (5 if weather in ['Storm','Fog'] else 0), scale=15)))
        if np.random.rand() < 0.02:
            delay += np.random.randint(60,300)
            
        # Save with the ORIGINAL column names for consistent renaming later
        rows.append({'airline':airline,'origin':origin,'destination':dest,
                      'sched_dep_hour':dep_hour,'sched_arr_hour':arr_hour,
                      'distance':distance,'day_of_week':day,'month':month,
                      'weather':weather,'delay_minutes':delay})
    df = pd.DataFrame(rows)
    df.to_csv(DATA_FILE, index=False)
    
    # Apply the rename immediately for use in the current session
    df.rename(columns=COLUMN_RENAMES, inplace=True)

# ***CRITICAL ADDITION TO PREVENT MAGIC COMMAND DISPLAY***
_ = True 

# ------------------- Model Training/Loading Logic -------------------
def train_and_save_model(df):
    """Trains a Linear Regression model and saves the model, encoders, scaler, AND performance."""
    st.info("💡 Training the Linear Regression model...")
    
    le_airline = LabelEncoder()
    le_origin = LabelEncoder()
    le_dest = LabelEncoder()
    le_weather = LabelEncoder()
    
    df_train = df.copy() 
    
    df_train['airline_enc'] = le_airline.fit_transform(df_train['airline'])
    df_train['origin_enc'] = le_origin.fit_transform(df_train['origin'])
    df_train['dest_enc'] = le_dest.fit_transform(df_train['dest'])
    df_train['weather_enc'] = le_weather.fit_transform(df_train['weather'])

    # CRITICAL: Features must use the NEW, RENAMED time columns: departure_time, arrival_time
    features = ['airline_enc', 'origin_enc', 'dest_enc', 'departure_time', 'arrival_time', 'distance', 'day_of_week', 'month', 'weather_enc']
    target = 'delay'
    X = df_train[features]
    y = df_train[target]

    numerical_cols = ['departure_time', 'arrival_time', 'distance', 'day_of_week', 'month']
    X_scaled = X.copy() 
    scaler = StandardScaler()
    X_scaled[numerical_cols] = scaler.fit_transform(X[numerical_cols])

    X_train, X_test, y_train, y_test = train_test_split(X_scaled, y, test_size=0.2, random_state=42)

    model = LinearRegression()
    model.fit(X_train, y_train) 
    
    # ----------------------------------------------------
    # CRITICAL FIX: Calculate and Save Performance Metrics
    # ----------------------------------------------------
    PERFORMANCE_FILE = 'outputs/model_performance.csv' # Output path for the Model Performance page
    
    y_pred = model.predict(X_test)
    
    # Calculate Regression Metrics
    mae = mean_absolute_error(y_test, y_pred)
    r2 = r2_score(y_test, y_pred)
    
    # Create and Save DataFrame
    performance_df = pd.DataFrame({
        'Metric': ['R-squared ($R^2$)', 'Mean Absolute Error (MAE)', 'Trained On Rows', 'Test Set Rows'],
        'Value': [f"{r2:.4f}", f"{mae:.2f}", len(X_train), len(X_test)]
    })
    
    performance_df.to_csv(PERFORMANCE_FILE, index=False) # SAVING THE FILE
    # ----------------------------------------------------

    # Save all model components
    joblib.dump(model, MODEL_FILE)
    joblib.dump(le_airline, 'models/le_airline.pkl')
    joblib.dump(le_origin, 'models/le_origin.pkl')
    joblib.dump(le_dest, 'models/le_dest.pkl')
    joblib.dump(le_weather, 'models/le_weather.pkl')
    joblib.dump(scaler, 'models/scaler.pkl')

    st.success("✅ Model training complete and files saved!")

# Force retraining if model is missing
if not os.path.exists(MODEL_FILE):
    train_and_save_model(df.copy())
# ------------------- End Model Training/Loading Logic -------------------


# ------------------- Sidebar Filters -------------------
st.sidebar.header("🔍 Filter Flights")
airline_filter = st.sidebar.multiselect("Airline", df['airline'].unique(), default=df['airline'].unique())
origin_filter = st.sidebar.multiselect("Origin", df['origin'].unique(), default=df['origin'].unique())
dest_filter = st.sidebar.multiselect("Destination", df['dest'].unique(), default=df['dest'].unique())
month_filter = st.sidebar.multiselect("Month", df['month'].unique(), default=df['month'].unique())
weather_filter = st.sidebar.multiselect("Weather", df['weather'].unique(), default=df['weather'].unique())

df_filtered = df[
    (df['airline'].isin(airline_filter)) &
    (df['origin'].isin(origin_filter)) &
    (df['dest'].isin(dest_filter)) &
    (df['month'].isin(month_filter)) &
    (df['weather'].isin(weather_filter))
]

# ------------------- Tabs -------------------
tab_overview, tab_analytics, tab_prediction, tab_report = st.tabs(
    ["Overview", "Analytics", "Predictions", "Reports"]
)

# ------------------- Overview Tab -------------------
with tab_overview:
    st.subheader("📊 Key Metrics")
    col1, col2, col3, col4 = st.columns(4)
    
    if len(df_filtered) > 0:
        avg_delay = int(df_filtered['delay'].mean())
        max_delay = df_filtered['delay'].max()
        min_delay = df_filtered['delay'].min()
        flights_count = len(df_filtered)

        col1.metric("Average Delay (min)", f"{avg_delay}", delta=f"{'+ High Delay' if avg_delay>30 else ''}")
        col2.metric("Maximum Delay (min)", f"{max_delay}")
        col3.metric("Minimum Delay (min)", f"{min_delay}")
        col4.metric("Flights Count", f"{flights_count}")

        st.subheader("📈 7-Day Average Delay Forecast (Synthetic)")
        future_dates = [datetime.today() + timedelta(days=i) for i in range(7)]
        forecast = [avg_delay + np.random.randint(-5,10) for _ in range(7)]
        fig_forecast = px.line(x=future_dates, y=forecast, markers=True, title="Forecasted Average Delay Next 7 Days")
        st.plotly_chart(fig_forecast, use_container_width=True)
    else:
        st.warning("No flights match the current filter selection.")

# ------------------- Analytics Tab (CORRECTED) -------------------
with tab_analytics:
    st.subheader("📊 Flight Analytics")
    
    if len(df_filtered) == 0:
        st.warning("No data to display. Adjust sidebar filters.")
    else:
        # Correction: Added expanded=True to all expanders
        with st.expander("⏰ Flight Delay Distribution", expanded=True):
            fig_delay = px.histogram(df_filtered, x="delay", nbins=40, color_discrete_sequence=["#FF6B6B"],
                                     title="Flight Delay Distribution",
                                     hover_data=['airline','origin','dest','distance','weather'])
            fig_delay.update_traces(marker_line_width=1, marker_line_color="white")
            st.plotly_chart(fig_delay, use_container_width=True)

        with st.expander("✈️ Top Airlines by Number of Flights", expanded=True):
            airline_counts = df_filtered['airline'].value_counts().reset_index()
            airline_counts.columns = ['airline','count']
            fig_airline = px.bar(airline_counts, x='airline', y='count', color='count',
                                 color_continuous_scale='Blues', text='count', title="Flights per Airline")
            st.plotly_chart(fig_airline, use_container_width=True)

        with st.expander("🌦️ Average Delay by Weather Condition", expanded=True):
            weather_delay = df_filtered.groupby('weather')['delay'].mean().reset_index()
            fig_weather = px.bar(weather_delay, x='weather', y='delay', color='delay',
                                 color_continuous_scale='Viridis', text='delay', title="Average Delay by Weather")
            st.plotly_chart(fig_weather, use_container_width=True)

        with st.expander("📏 Delay vs Distance by Weather", expanded=True):
            fig_scatter = px.scatter(df_filtered, x='distance', y='delay', color='weather', size='delay',
                                     hover_data=['airline','origin','dest'], title="Delay vs Distance")
            st.plotly_chart(fig_scatter, use_container_width=True)

        with st.expander("🗺️ Flight Routes Map", expanded=True):
            airport_coords = {
                'BLR': (12.9716,77.5946), 'DEL': (28.7041,77.1025), 'BOM': (19.0760,72.8777),
                'MAA': (13.0827,80.2707), 'COK': (9.9312,76.2673), 'HYD': (17.3850,78.4867),
                'CCU': (22.5726,88.3639), 'AMD': (23.0225,72.5714)
            }
            df_map = df_filtered.copy()
            df_map['origin_lat'] = df_map['origin'].map(lambda x: airport_coords.get(x, (0,0))[0])
            df_map['origin_lon'] = df_map['origin'].map(lambda x: airport_coords.get(x, (0,0))[1])
            df_map['dest_lat'] = df_map['dest'].map(lambda x: airport_coords.get(x, (0,0))[0])
            df_map['dest_lon'] = df_map['dest'].map(lambda x: airport_coords.get(x, (0,0))[1])

            fig_map = go.Figure()
            for i,row in df_map.iterrows():
                fig_map.add_trace(go.Scattergeo(
                    locationmode='ISO-3',
                    lon=[row['origin_lon'], row['dest_lon']],
                    lat=[row['origin_lat'], row['dest_lat']],
                    mode='lines+markers',
                    line=dict(width=2, color='red' if row['delay']>60 else 'green'),
                    marker=dict(size=4),
                    hoverinfo='text',
                    text=f"{row['airline']}: {row['origin']} → {row['dest']} | Delay: {row['delay']} min"
                ))
            fig_map.update_layout(
                geo=dict(scope='asia', projection_type='natural earth', center={'lon': 78, 'lat': 20},
                          lataxis_range=[5, 35], lonaxis_range=[65, 95]), 
                title="Flight Routes with Delay Highlighted"
            )
            st.plotly_chart(fig_map, use_container_width=True)

# ------------------- Predictions Tab -------------------
with tab_prediction:
    st.subheader("🔮 Predict Flight Delay")
    st.markdown("Single flight or batch CSV prediction")

    # Single prediction
    with st.form("predict_form"):
        airline_in = st.selectbox("Airline", df['airline'].unique())
        origin_in = st.selectbox("Origin", df['origin'].unique())
        dest_in = st.selectbox("Destination", df['dest'].unique())
        dep_hour_in = st.slider("Scheduled Departure Hour", 0, 23, 10)
        arr_hour_in = st.slider("Scheduled Arrival Hour", 0, 23, 12)
        distance_in = st.number_input("Distance (km)", 50, 5000, 500)
        day_in = st.slider("Day of Week (0=Mon)", 0, 6, 1)
        month_in = st.slider("Month", 1, 12, 1)
        weather_in = st.selectbox("Weather", df['weather'].unique())
        submit_btn = st.form_submit_button("Predict Delay")

        if submit_btn:
            if os.path.exists(MODEL_FILE):
                try:
                    model = joblib.load(MODEL_FILE)
                    le_airline = joblib.load('models/le_airline.pkl')
                    le_origin = joblib.load('models/le_origin.pkl')
                    le_dest = joblib.load('models/le_dest.pkl')
                    le_weather = joblib.load('models/le_weather.pkl')
                    scaler = joblib.load('models/scaler.pkl')

                    features = ['airline_enc', 'origin_enc', 'dest_enc', 'departure_time', 'arrival_time', 'distance', 'day_of_week', 'month', 'weather_enc']

                    X_input = pd.DataFrame({
                        'airline_enc':[le_airline.transform([airline_in])[0]],
                        'origin_enc':[le_origin.transform([origin_in])[0]],
                        'dest_enc':[le_dest.transform([dest_in])[0]],
                        'departure_time':[dep_hour_in],
                        'arrival_time':[arr_hour_in],
                        'distance':[distance_in],
                        'day_of_week':[day_in],
                        'month':[month_in],
                        'weather_enc':[le_weather.transform([weather_in])[0]]
                    })
                    
                    cols_to_scale = ['departure_time','arrival_time','distance','day_of_week','month']
                    X_input[cols_to_scale] = scaler.transform(X_input[cols_to_scale])
                    
                    X_predict = X_input[features] 

                    pred_delay = model.predict(X_predict)[0]
                    predicted_delay_minutes = max(0, int(round(pred_delay)))
                    
                    st.success(f"Estimated Flight Delay: **{predicted_delay_minutes} minutes**")
                    if predicted_delay_minutes >= 30:
                        st.balloons()
                        st.warning("This flight is predicted to have a significant delay (30+ minutes).")
                        
                except Exception as e:
                    if "unseen label" in str(e) or "not found" in str(e):
                        st.error("Error: The selected categorical value was not present in the training data.")
                    else:
                        st.error(f"Error during prediction: {str(e)}")
            else:
                st.error("⚠️ Trained model not found. Please refresh the app to run the training!")

    # Batch prediction
    st.subheader("📂 Batch Prediction via CSV")
    uploaded_file = st.file_uploader("Upload CSV", type=["csv"])
    if uploaded_file:
        batch_df = pd.read_csv(uploaded_file)
        st.dataframe(batch_df.head())
        if st.button("Predict Batch Delays"):
            try:
                model = joblib.load(MODEL_FILE)
                le_airline = joblib.load('models/le_airline.pkl')
                le_origin = joblib.load('models/le_origin.pkl')
                le_dest = joblib.load('models/le_dest.pkl')
                le_weather = joblib.load('models/le_weather.pkl')
                scaler = joblib.load('models/scaler.pkl')
                
                batch_df.rename(columns=COLUMN_RENAMES, inplace=True, errors='ignore')

                batch_df['airline_enc'] = le_airline.transform(batch_df['airline'])
                batch_df['origin_enc'] = le_origin.transform(batch_df['origin'])
                batch_df['dest_enc'] = le_dest.transform(batch_df['dest'])
                batch_df['weather_enc'] = le_weather.transform(batch_df['weather'])
                cols_to_scale = ['departure_time','arrival_time','distance','day_of_week','month']
                
                batch_df_scaled = batch_df.copy()
                batch_df_scaled[cols_to_scale] = scaler.transform(batch_df_scaled[cols_to_scale])
                
                features_to_predict = ['airline_enc','origin_enc','dest_enc','departure_time','arrival_time','distance','day_of_week','month','weather_enc']

                batch_df['predicted_delay'] = model.predict(batch_df_scaled[features_to_predict])
                
                batch_df['predicted_delay'] = batch_df['predicted_delay'].apply(lambda x: max(0, int(round(x))))
                
                st.success("Batch prediction completed!")
                st.dataframe(batch_df.head())

                csv_buffer = BytesIO()
                batch_df.to_csv(csv_buffer, index=False)
                st.download_button("Download Predictions CSV", data=csv_buffer.getvalue(), file_name="batch_predictions.csv")
            except Exception as e:
                st.error(f"Error during batch prediction. Check your CSV column names and ensure they match the model's training data. Error: {str(e)}")

# ------------------- Reports Tab -------------------
with tab_report:
    st.subheader("📄 Download DOCX Report")
    
    if len(df_filtered) == 0:
        st.warning("No data to generate report from. Adjust sidebar filters.")
    else:
        airline_counts = df_filtered['airline'].value_counts().reset_index()
        airline_counts.columns = ['airline','count']
        fig_airline = px.bar(airline_counts, x='airline', y='count', color='count', color_continuous_scale='Blues', text='count', title="Flights per Airline")

        weather_delay = df_filtered.groupby('weather')['delay'].mean().reset_index()
        fig_weather = px.bar(weather_delay, x='weather', y='delay', color='delay', color_continuous_scale='Viridis', text='delay', title="Average Delay by Weather")
        
        fig_delay = px.histogram(df_filtered, x="delay", nbins=40, color_discrete_sequence=["#FF6B6B"], title="Flight Delay Distribution")
        
        fig_scatter = px.scatter(df_filtered, x='distance', y='delay', color='weather', size='delay', title="Delay vs Distance")


        if st.button("Generate DOCX Report"):
            try:
                doc = Document()
                doc.add_heading("Flight Delay Prediction Dashboard Report", 0)
                doc.add_paragraph(f"Report Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
                doc.add_paragraph("---")
                
                doc.add_heading("Key Metrics (Filtered Data)", level=1)
                doc.add_paragraph(f"Filtered Flights Count: {len(df_filtered)}")
                doc.add_paragraph(f"Average Delay: {df_filtered['delay'].mean():.2f} minutes")
                doc.add_paragraph(f"Maximum Delay: {df_filtered['delay'].max()} minutes")
                doc.add_paragraph("---")

                for i, fig in enumerate([fig_delay, fig_airline, fig_weather, fig_scatter]):
                    out_path = f"outputs/chart_{uuid.uuid4().hex}.png"
                    fig.write_image(out_path)
                    doc.add_picture(out_path, width=Inches(6))
                    os.remove(out_path) 
                    doc.add_paragraph("\n")

                report_buffer = BytesIO()
                doc.save(report_buffer)
                report_buffer.seek(0)
                
                st.download_button(
                    label="Download DOCX Report",
                    data=report_buffer,
                    file_name="Flight_Delay_Report.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
                st.success(f"✅ Report generated successfully!")
            except Exception as e:
                st.error(f"Error generating report: {str(e)}. Ensure you have 'kaleido' installed for image export (`pip install python-docx plotly kaleido`).")

st.caption("© 2025 Nirmal Raj")